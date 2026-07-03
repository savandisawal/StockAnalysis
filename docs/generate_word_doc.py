"""
Convert FUNCTIONAL_GUIDE.md to a formatted Word document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD_PATH = "FUNCTIONAL_GUIDE.md"
OUT_PATH = "Stock_Quant_App_Functional_Guide.docx"

# ── colour palette ─────────────────────────────────────────────────────────
BRAND_BLUE  = RGBColor(0x1F, 0x4E, 0x79)   # dark navy  – h1/h2
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)   # mid blue   – h3
ACCENT_TEAL = RGBColor(0x17, 0x6A, 0x6A)   # teal       – h4
TABLE_HDR   = RGBColor(0x2E, 0x75, 0xB6)   # table header bg
LIGHT_GREY  = RGBColor(0xF2, 0xF2, 0xF2)   # alt table row
CODE_BG     = RGBColor(0xF5, 0xF5, 0xF5)


def set_cell_bg(cell, colour_hex: str):
    """Fill a table cell background with a hex colour string like '2E75B6'."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  colour_hex)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = BRAND_BLUE
        run.font.size = Pt(20)
    elif level == 2:
        run.font.color.rgb = BRAND_BLUE
        run.font.size = Pt(16)
    elif level == 3:
        run.font.color.rgb = ACCENT_BLUE
        run.font.size = Pt(13)
    elif level == 4:
        run.font.color.rgb = ACCENT_TEAL
        run.font.size = Pt(12)
    return p


def add_normal(doc: Document, text: str):
    """Add a paragraph, rendering **bold** inline markup."""
    p = doc.add_paragraph()
    _add_inline(p, text)
    return p


def _add_inline(para, text: str):
    """Split text on **bold** markers and add runs accordingly."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def add_code_block(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
    # light grey shading on the paragraph
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F5F5F5")
    pPr.append(shd)
    return p


def add_md_table(doc: Document, rows: list[str]):
    """Parse a Markdown table (list of raw lines) and add a formatted Word table."""
    # strip leading/trailing pipes and split
    cleaned = []
    for r in rows:
        r = r.strip().strip("|")
        cleaned.append([c.strip() for c in r.split("|")])

    # row 1 = header, row 2 = separator (skip), rest = data
    header = cleaned[0]
    data   = [r for r in cleaned[2:] if r]

    col_count = len(header)
    table = doc.add_table(rows=1 + len(data), cols=col_count)
    table.style = "Table Grid"

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(header):
        cell = hdr_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_bg(cell, "2E75B6")

    # data rows
    for r_idx, row in enumerate(data):
        tr = table.rows[r_idx + 1]
        fill = "F2F2F2" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = ""
            _add_inline(cell.paragraphs[0], val)
            cell.paragraphs[0].runs and setattr(
                cell.paragraphs[0].runs[0].font, "size", Pt(10)
            )
            set_cell_bg(cell, fill)

    doc.add_paragraph()   # spacing after table


def add_bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    _add_inline(p, text)
    return p


# ── Cover page ─────────────────────────────────────────────────────────────

def build_cover(doc: Document):
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Stock Quant App")
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = BRAND_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Functional Guide")
    run2.font.size  = Pt(18)
    run2.font.color.rgb = ACCENT_BLUE

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = note.add_run("NSE India — Next-Day Price Range Prediction System")
    run3.font.size = Pt(12)
    run3.font.italic = True
    run3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_page_break()


# ── Main parser ────────────────────────────────────────────────────────────

def parse_and_build(doc: Document, md_path: str):
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_code = False
    code_buf: list[str] = []
    md_table_buf: list[str] = []

    def flush_table():
        nonlocal md_table_buf
        if md_table_buf:
            add_md_table(doc, md_table_buf)
            md_table_buf = []

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # ── code block ──────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                add_code_block(doc, "\n".join(code_buf))
                in_code = False
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── markdown table ──────────────────────────────────────────────
        if "|" in line and line.strip().startswith("|"):
            md_table_buf.append(line)
            i += 1
            continue
        else:
            flush_table()

        stripped = line.strip()

        # ── horizontal rule ─────────────────────────────────────────────
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # ── headings ────────────────────────────────────────────────────
        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:], 4)
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:], 3)
        elif stripped.startswith("## "):
            add_heading(doc, stripped[3:], 2)
        elif stripped.startswith("# "):
            add_heading(doc, stripped[2:], 1)

        # ── bullet points ───────────────────────────────────────────────
        elif stripped.startswith("- ") or stripped.startswith("* "):
            add_bullet(doc, stripped[2:], level=0)
        elif re.match(r"^\d+\. ", stripped):
            add_bullet(doc, re.sub(r"^\d+\. ", "", stripped), level=0)

        # ── blank line ──────────────────────────────────────────────────
        elif stripped == "":
            pass   # skip blank lines (Word already has spacing)

        # ── normal paragraph ────────────────────────────────────────────
        else:
            add_normal(doc, stripped)

        i += 1

    flush_table()


# ── Entry point ────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import os
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path  = os.path.join(script_dir, MD_PATH)
    out_path = os.path.join(script_dir, OUT_PATH)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    build_cover(doc)
    parse_and_build(doc, md_path)

    doc.save(out_path)
    print(f"Saved: {out_path}")
